import pytesseract, time, fitz
from docx import Document
from openpyxl import load_workbook
from pathlib import Path
from typing import Optional, Dict, Any
from PIL import Image

from utils import convert_pptx_to_pdf, convert_doc_to_pdf
from ia import OpenAIInference, GeminiInferenceForImages

class OfficeDocumentExtractor:
    """
    Una clase unificada para extraer contenido de documentos de Microsoft Office.
    Maneja documentos Word (.docx), Excel (.xlsx) y PowerPoint (.pptx).
    """
    def __init__(self):
        self.ia_inference = OpenAIInference()
        self.ia_inference_for_images = GeminiInferenceForImages()

    def extract_docx(self, file_path: str) -> str:
        """
        Extrae el contenido de un archivo Word (.docx).
        Esta función obtiene tanto el texto de los párrafos como el contenido de las tablas.
        """
        doc = Document(file_path)
        content = []

        # Extraemos el texto de cada párrafo del documento
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Solo añadimos párrafos no vacíos
                content.append(paragraph.text)

        # Extraemos el contenido de las tablas
        for table in doc.tables:
            for row in table.rows:
                # Unimos las celdas de cada fila con tabulaciones
                row_content = '\t'.join(cell.text for cell in row.cells)
                if row_content.strip():  # Solo añadimos filas no vacías
                    content.append(row_content)

        return '\n'.join(content)

    def extract_xlsx(self, file_path: str) -> Dict[str, list]:
        """
        Extrae el contenido de un archivo Excel (.xlsx).
        Devuelve un diccionario donde cada clave es el nombre de una hoja
        y el valor es una lista de filas con su contenido.
        """
        workbook = load_workbook(file_path, data_only=True)
        content = {}

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_content = []

            # Iteramos por cada fila de la hoja
            for row in sheet.iter_rows(values_only=True):
                # Filtramos las celdas vacías y convertimos todo a string
                row_content = [str(cell) if cell is not None else '' for cell in row]
                if any(cell.strip() for cell in row_content):  # Solo añadimos filas no vacías
                    sheet_content.append(row_content)

            content[sheet_name] = sheet_content

        return content

    def extract_pdf(self, file_path):
        file = fitz.open(file_path)
        for page in file:
            try:
                # Convertimos la página a imagen
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                # Realizamos OCR
                text = pytesseract.image_to_string(img)
                time.sleep(0.1)
                # print(text)
                # user_data = get_inference_for_pdf_open_ai(text)
                return text
            except Exception as e:
                print(f"Error en el procesamiento del pdf usando OCR: {str(e)}")
                return None
        file.close()
    
    def extract_content(self, file_path: str) -> Optional[Any]:
        """
        Método principal que determina el tipo de archivo y llama al método
        apropiado para extraer su contenido.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"El archivo {file_path} no existe")

        try:
            if file_path.suffix.lower() == '.pdf':
                content_doc = self.extract_pdf(str(file_path))
                print('AAAAAAAAAAAAA',file_path)
                content = self.ia_inference.get_inference(content_doc, file_path)
                return content
            elif file_path.suffix.lower() == '.docx' or file_path.suffix.lower() == '.doc':
                try:
                    pdf_path = convert_doc_to_pdf(file_path)
                    content_pdf = self.extract_pdf(pdf_path)
                    content = self.ia_inference.get_inference(content_pdf)
                except Exception as e:
                    print(f"Error procesando el PDF: {e}")
                    raise
                else:
                    pdf_path.unlink()
                    return content                   
            elif file_path.suffix.lower() == '.xlsx':
                content_doc = self.extract_xlsx(str(file_path))
                content = self.ia_inference.get_inference(str(content_doc))
                return content
            elif file_path.suffix.lower() == '.pptx':
                try:
                    pdf_path = convert_pptx_to_pdf(file_path)
                    content_pdf = self.extract_pdf(pdf_path)
                    content = self.ia_inference.get_inference(content_pdf)
                except Exception as e:
                    print(f"Error procesando el PDF: {e}")
                    raise
                else:
                    pdf_path.unlink()
                    return content

            elif file_path.suffix.lower() == '.jpg' or file_path.suffix.lower() == '.JPG' or file_path.suffix.lower() == '.jpeg' or file_path.suffix.lower() == '.png' or file_path.suffix.lower() == '.PNG':
                content = self.ia_inference_for_images.get_inference(file_path)
                return content
            else:
                raise ValueError(f"Formato de archivo no soportado: {file_path.suffix}")

        except Exception as e:
            raise

if __name__ == "__main__":
    extractor = OfficeDocumentExtractor()
    xlsx_content = extractor.extract_content("/home/desarrollo/Documents/wc/processing-certificates/certificates/1/23002896_clarissa_quintana_ramos_certificado_reanimacion.docx")
    print (xlsx_content)