"""
Sistema de extracción y reporte de certificaciones médicas
Genera reportes Excel con información completa de certificados incluyendo transcripción
NUEVA FUNCIONALIDAD: Manejo de PDFs multipágina con separación automática
"""

import os
import logging
import json
import base64
import time
import fitz
import pytesseract
from abc import ABC, abstractmethod
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
from PIL import Image
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
from openai import OpenAI, RateLimitError, APIConnectionError, Timeout, OpenAIError
import anthropic
from datetime import datetime
import shutil

# Cargar variables de entorno
load_dotenv()

# =============================================================================
# CONFIGURACIÓN DE MODELO POR DEFECTO
# =============================================================================
# Cambia esta variable para usar el modelo que prefieras:
# Opciones: "openai", "anthropic", "gemini"
DEFAULT_MODEL = "gemini"  # ← CAMBIA AQUÍ EL MODELO POR DEFECTO
# =============================================================================

# Configuración de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

logger.info(f"🤖 Modelo configurado por defecto: {DEFAULT_MODEL.upper()}")

class ArchivoProhibido(Enum):
    """Enum para extensiones de archivos prohibidos"""
    INI = '.ini'
    EXE = '.exe'
    HTML = '.html'
    BAK = '.bak'
    TMP = '.tmp'

class InferenceError(Exception):
    """Clase personalizada para errores de inferencia"""
    pass

class ResponseParsingError(Exception):
    """Clase personalizada para errores de parsing de respuesta"""
    pass

class BaseInferenceReport(ABC):
    """
    Clase base abstracta para extracción de información completa de certificados.
    Extrae todos los campos necesarios para el reporte completo.
    """
    def __init__(self):
        self.prompt = """Tu tarea es extraer información específica y completa de certificados escaneados, siendo especialmente flexible con variaciones en el texto debido a OCR.

INSTRUCCIONES DE EXTRACCIÓN - TODOS LOS CAMPOS OBLIGATORIOS:

1. NOMBRE DEL CERTIFICADO/CURSO:
   - Busca el nombre completo del certificado o curso
   - Palabras clave: 'CERTIFICA', 'CURSO DE', 'CERTIFICADO EN', 'DIPLOMA DE', 'ENTRENAMIENTO EN'
   - Ejemplo: "SOPORTE VITAL BASICO Y AVANZADO", "ACLS", "REANIMACION PEDIATRICA"
   - Si no se encuentra: null

2. NOMBRE COMPLETO DEL PARTICIPANTE:
   - Busca el nombre después de: 'CERTIFICA QUE', 'OTORGADO A', 'NOMBRE:', 'PARTICIPANTE:'
   - Incluye nombres y apellidos completos
   - Limpia caracteres especiales del OCR
   - Si no se encuentra: null

3. IDENTIFICACIÓN:
   - Busca números precedidos por: 'C.C.', 'CC', 'Cédula', 'C.I.', 'DNI'
   - Extrae solo números, sin puntos ni espacios
   - Si no se encuentra: null

4. INSTITUCIÓN:
   - Busca nombre de la institución emisora
   - Palabras clave: 'INSTITUTO', 'FUNDACIÓN', 'UNIVERSIDAD', 'CENTRO', 'ACADEMIA'
   - Incluye nombre completo de la organización
   - Si no se encuentra: null

5. CIUDAD:
   - Busca ubicación geográfica
   - Palabras clave: 'CIUDAD DE', 'EN', ubicaciones después de fechas
   - Ejemplos: "Bogotá", "Barranquilla", "Medellín"
   - Si no se encuentra: null

6. FECHA DE EMISIÓN:
   - Busca patrones de fecha con palabras clave: 'Realizado', 'Expedido', 'Emitido', 'Fecha'
   - Formatos: "DD de MM de YYYY", "DD/MM/YYYY", "MM-DD-YYYY"
   - Convierte a formato ISO: "YYYY-MM-DDTHH:MM:SS.sssZ"
   - Si no se encuentra: null

7. FECHA DE EXPIRACIÓN:
   - Busca: 'válido hasta', 'vigencia', 'vence', 'expira'
   - Si encuentra vigencia en años/meses/días: calcula desde fecha de emisión
   - Convierte a formato ISO: "YYYY-MM-DDTHH:MM:SS.sssZ"
   - Si no se encuentra: null

8. INTENSIDAD HORARIA:
   - Busca números seguidos de: 'horas', 'hrs', 'h'
   - Extrae solo el número
   - Ejemplos: "48 horas" → 48, "16 hrs" → 16
   - Si no se encuentra: null

9. DIRIGIDO A (PERFIL PROFESIONAL):
   - Busca: 'Para:', 'Dirigido a:', profesiones mencionadas
   - Ejemplos: "Médico", "Enfermero", "Paramédico", "Personal de salud"
   - Si no se encuentra: null

10. ÁREA DE ESPECIALIZACIÓN:
    - Identifica el área médica basada en el contenido
    - Ejemplos: "Cardiología", "Pediatría", "Emergencias", "Trasplantes", "Cirugía"
    - Inferir del nombre del curso si no está explícito
    - Si no se encuentra: null

CAMPOS ADICIONALES OPCIONALES:

11. NIVEL DEL CURSO:
    - Busca: 'Básico', 'Avanzado', 'Intermedio', 'Especializado'
    - Inferir del nombre si no está explícito
    - Si no se encuentra: null

12. LINEAMIENTOS/ESTÁNDARES:
    - Busca: 'AHA', 'ILCOR', 'ACC', 'ESC', menciones de sociedades científicas
    - Si no se encuentra: null

13. INSTRUCTOR PRINCIPAL:
    - Busca nombres en secciones de firmas o instructores
    - Si no se encuentra: null

14. NIT DE INSTITUCIÓN:
    - Busca números precedidos por 'NIT', 'NIT:', 'Nit'
    - Si no se encuentra: null

FORMATO DE SALIDA OBLIGATORIO:
{
    "certificate_name": string | null,
    "participant_name": string | null,
    "identification": string | null,
    "institution": string | null,
    "city": string | null,
    "issue_date": "YYYY-MM-DDTHH:MM:SS.sssZ" | null,
    "expiration_date": "YYYY-MM-DDTHH:MM:SS.sssZ" | null,
    "hours": number | null,
    "target_audience": string | null,
    "specialization_area": string | null,
    "level": string | null,
    "guidelines": string | null,
    "instructor": string | null,
    "institution_nit": string | null
}

REGLAS IMPORTANTES:
- TODOS los campos deben aparecer en la respuesta, aunque sean null
- Sé muy tolerante con errores de OCR
- Busca patrones aproximados cuando el texto esté distorsionado
- NO agregues campos adicionales no solicitados
- NO incluyas explicaciones, solo el JSON
- Si un campo no se puede determinar con certeza, asigna null"""

    def parse_response(self, response_text: str) -> dict:
        """
        Procesa y valida las respuestas JSON asegurando que todos los campos estén presentes.
        """
        try:
            # Extraer JSON de la respuesta
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}')
            
            if start_idx == -1 or end_idx == -1:
                raise ResponseParsingError("No se encontró una estructura JSON válida en la respuesta")
            
            json_str = response_text[start_idx:end_idx + 1]
            data = json.loads(json_str)
            
            # Campos obligatorios con valores por defecto
            required_fields = {
                'certificate_name': None,
                'participant_name': None,
                'identification': None,
                'institution': None,
                'city': None,
                'issue_date': None,
                'expiration_date': None,
                'hours': None,
                'target_audience': None,
                'specialization_area': None,
                'level': None,
                'guidelines': None,
                'instructor': None,
                'institution_nit': None,
                'message_error': None
            }
            
            # Asegurar que todos los campos estén presentes
            for field in required_fields:
                if field not in data:
                    data[field] = required_fields[field]
            
            return data
        
        except json.JSONDecodeError as e:
            logger.error(f"Error al decodificar JSON: {str(e)}")
            logger.error(f"Texto recibido: {response_text}")
            raise ResponseParsingError(f"Error al decodificar JSON: {str(e)}")
        except Exception as e:
            logger.error(f"Error inesperado al procesar respuesta: {str(e)}")
            raise ResponseParsingError(f"Error inesperado: {str(e)}")

    @abstractmethod
    def get_inference(self, content):
        """Método abstracto para obtener inferencias"""
        pass

class OpenAIInferenceReport(BaseInferenceReport):
    """Clase para manejar inferencia con OpenAI para reportes completos"""
    
    def __init__(self):
        super().__init__()
        self.api_key = os.getenv('API_OPENAI')
        self.client = OpenAI(api_key=self.api_key)

    def get_inference(self, content_certificate, path_pdf=None):
        try:
            completion = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": self.prompt},
                    {"role": "user", "content": content_certificate}
                ],
                temperature=0.1
            )
            
            response_text = completion.choices[0].message.content
            if not response_text:
                raise Exception("La API no devolvió ninguna respuesta")
            
            certificate_info = self.parse_response(response_text)
            return certificate_info
            
        except (RateLimitError, APIConnectionError, Timeout, OpenAIError) as e:
            logger.error(f"Error en API de OpenAI: {str(e)}")
            raise InferenceError(f"Error en OpenAI: {str(e)}")
        except Exception as e:
            logger.error(f"Error inesperado: {str(e)}")
            raise InferenceError(f"Error inesperado: {str(e)}")

class AntropicInferenceReport(BaseInferenceReport):
    """Clase para manejar inferencia con Claude para reportes completos"""
    
    def __init__(self):
        super().__init__()
        self.api_key = os.getenv('ANTHROPIC_API_KEY')
        self.client = anthropic.Anthropic(api_key=self.api_key)
        self.model = "claude-3-5-sonnet-20241022"

    def read_pdf(self, path_pdf):
        try:
            with open(path_pdf, "rb") as f:
                pdf_data = base64.b64encode(f.read()).decode("utf-8")
            return pdf_data
        except FileNotFoundError:
            raise InferenceError("El archivo PDF no se encontró")
    
    def get_inference(self, content_certificate=None, path_pdf=None):
        try:
            if path_pdf:
                # Usar PDF directamente con Claude
                content_pdf = self.read_pdf(path_pdf)
                message = self.client.messages.create(
                    model=self.model,
                    max_tokens=2048,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": content_pdf
                                    }
                                },
                                {
                                    "type": "text",
                                    "text": self.prompt
                                }
                            ]
                        }
                    ]
                )
            else:
                # Usar texto extraído
                message = self.client.messages.create(
                    model=self.model,
                    max_tokens=2048,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": f"{self.prompt}\n\nTexto del certificado:\n{content_certificate}"
                                }
                            ]
                        }
                    ]
                )
            
            text_response = message.content[0]
            response_dict = json.loads(text_response.text)
            response_dict['message_error'] = None
            
            return response_dict

        except anthropic.APIError as e:
            error_message = f"Error en la API de Anthropic: {str(e)}"
            logger.error(error_message)
            raise InferenceError(error_message)
        except Exception as e:
            error_message = f"Error inesperado en Claude: {str(e)}"
            logger.error(error_message)
            raise InferenceError(error_message)

class GeminiInferenceReport(BaseInferenceReport):
    """Clase para manejar inferencia con Gemini para documentos e imágenes"""
    
    def __init__(self):
        super().__init__()
        self.api_key = os.getenv('GEMINI_API')
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')
        
    def get_inference(self, content_certificate=None, path_file=None):
        try:
            if path_file and Path(path_file).suffix.lower() in ['.jpg', '.jpeg', '.png']:
                # Procesar imagen
                image = Image.open(path_file)
                response = self.model.generate_content([self.prompt, image])
                response.resolve()
            else:
                # Procesar texto extraído
                response = self.model.generate_content(f"{self.prompt}\n\nTexto del certificado:\n{content_certificate}")
                response.resolve()

            if response.text:
                return self.parse_response(response.text)
            else:
                raise InferenceError("No se generó respuesta de texto")

        except FileNotFoundError:
            raise InferenceError("El archivo no se encontró")
        except Exception as e:
            raise InferenceError(f"Error en Gemini: {str(e)}")

class ModelFactory:
    """Factory para crear instancias de modelos según la configuración"""
    
    @staticmethod
    def create_model(model_name: str):
        """Crea una instancia del modelo especificado"""
        model_name = model_name.lower()
        
        if model_name == "openai":
            logger.info("🔷 Usando modelo OpenAI (GPT-3.5-turbo)")
            return OpenAIInferenceReport()
        elif model_name == "anthropic":
            logger.info("🟡 Usando modelo Anthropic (Claude 3.5 Sonnet)")
            return AntropicInferenceReport()
        elif model_name == "gemini":
            logger.info("🔵 Usando modelo Google Gemini (1.5 Flash)")
            return GeminiInferenceReport()
        else:
            logger.warning(f"Modelo '{model_name}' no reconocido. Usando OpenAI por defecto.")
            return OpenAIInferenceReport()

class PDFSplitter:
    """🆕 NUEVA CLASE: Maneja la separación de PDFs multipágina"""
    
    def __init__(self):
        pass
    
    def get_pdf_page_count(self, file_path: str) -> int:
        """Obtiene el número de páginas de un PDF"""
        try:
            pdf_document = fitz.open(file_path)
            page_count = pdf_document.page_count
            pdf_document.close()
            return page_count
        except Exception as e:
            logger.error(f"Error obteniendo número de páginas de {file_path}: {str(e)}")
            return 1  # Asumir 1 página si hay error
    
    def split_pdf_pages(self, pdf_path: str, output_directory: str) -> Tuple[bool, list]:
        """
        Separa un PDF en páginas individuales
        
        Args:
            pdf_path: Ruta del PDF original
            output_directory: Directorio donde guardar las páginas separadas
            
        Returns:
            Tuple[bool, list]: (éxito, lista_de_archivos_creados)
        """
        try:
            # Crear directorio de salida si no existe
            os.makedirs(output_directory, exist_ok=True)
            
            # Abrir PDF original
            pdf_document = fitz.open(pdf_path)
            
            # Obtener nombre base del archivo
            base_name = Path(pdf_path).stem
            created_files = []
            
            logger.info(f"📄 Separando PDF '{base_name}' en {pdf_document.page_count} páginas...")
            
            # Separar cada página
            for page_num in range(pdf_document.page_count):
                # Crear nuevo PDF con una sola página
                new_pdf = fitz.open()
                new_pdf.insert_pdf(pdf_document, from_page=page_num, to_page=page_num)
                
                # Nombre del archivo de la página
                page_filename = f"{base_name}_pagina_{page_num + 1:02d}.pdf"
                page_path = os.path.join(output_directory, page_filename)
                
                # Guardar página
                new_pdf.save(page_path)
                new_pdf.close()
                
                created_files.append(page_path)
                logger.info(f"   ✅ Página {page_num + 1} guardada como: {page_filename}")
            
            pdf_document.close()
            
            logger.info(f"🎉 PDF separado exitosamente en {len(created_files)} archivos")
            return True, created_files
            
        except Exception as e:
            logger.error(f"❌ Error separando PDF {pdf_path}: {str(e)}")
            return False, []

class DocumentExtractorReport:
    """Extractor unificado para documentos con reporte completo - MEJORADO para manejar PDFs multipágina"""
    
    def __init__(self):
        # Crear modelo principal según configuración
        self.ia_inference = ModelFactory.create_model(DEFAULT_MODEL)
        
        # Crear modelos de respaldo para casos específicos
        self.fallback_models = {
            "openai": OpenAIInferenceReport(),
            "anthropic": AntropicInferenceReport(),
            "gemini": GeminiInferenceReport()
        }
        
        # Remover el modelo principal de los fallbacks
        if DEFAULT_MODEL in self.fallback_models:
            del self.fallback_models[DEFAULT_MODEL]
        
        # 🆕 NUEVA FUNCIONALIDAD: Inicializar splitter de PDFs
        self.pdf_splitter = PDFSplitter()

    def extract_pdf_content(self, file_path):
        """Extrae contenido de PDF usando OCR"""
        try:
            file = fitz.open(file_path)
            full_text = ""
            
            for page in file:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(img)
                    full_text += text + "\n"
                    time.sleep(0.1)
                except Exception as e:
                    logger.error(f"Error procesando página: {str(e)}")
                    continue
            
            file.close()
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"Error extrayendo PDF: {str(e)}")
            raise InferenceError(f"Error extrayendo PDF: {str(e)}")

    def extract_docx_content(self, file_path):
        """Extrae contenido de documento Word"""
        try:
            doc = Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_content = '\t'.join(cell.text for cell in row.cells)
                    if row_content.strip():
                        content.append(row_content)

            return '\n'.join(content)
            
        except Exception as e:
            logger.error(f"Error extrayendo DOCX: {str(e)}")
            raise InferenceError(f"Error extrayendo DOCX: {str(e)}")

    def should_split_pdf(self, file_path: str) -> bool:
        """🆕 NUEVA FUNCIÓN: Determina si un PDF debe ser separado"""
        if Path(file_path).suffix.lower() != '.pdf':
            return False
        
        page_count = self.pdf_splitter.get_pdf_page_count(file_path)
        return page_count > 1

    def extract_content(self, file_path: str) -> tuple:
        """
        Extrae contenido y obtiene información del certificado
        🆕 MEJORADO: Ahora maneja separación automática de PDFs multipágina
        
        Retorna: (info_certificado, transcripcion_completa)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"El archivo {file_path} no existe")

        try:
            transcription = ""
            certificate_info = {}
            
            if file_path.suffix.lower() == '.pdf':
                if DEFAULT_MODEL == "anthropic":
                    # Claude puede procesar PDFs directamente
                    certificate_info = self.ia_inference.get_inference(path_pdf=str(file_path))
                    transcription = "PDF procesado directamente por Claude"
                else:
                    # Para otros modelos, extraer texto primero
                    transcription = self.extract_pdf_content(str(file_path))
                    certificate_info = self.ia_inference.get_inference(transcription, str(file_path))
                
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                transcription = self.extract_docx_content(str(file_path))
                certificate_info = self.ia_inference.get_inference(transcription)
                
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
                if DEFAULT_MODEL == "gemini":
                    # Gemini puede procesar imágenes directamente
                    certificate_info = self.ia_inference.get_inference(path_file=str(file_path))
                    transcription = f"Imagen procesada directamente por Gemini: {file_path.name}"
                else:
                    # Para otros modelos, usar Gemini como fallback para imágenes
                    certificate_info = self.fallback_models.get("gemini", GeminiInferenceReport()).get_inference(path_file=str(file_path))
                    transcription = f"Imagen procesada por Gemini (fallback): {file_path.name}"
                
            else:
                raise ValueError(f"Formato no soportado: {file_path.suffix}")

            return certificate_info, transcription

        except Exception as e:
            logger.error(f"Error procesando {file_path}: {str(e)}")
            
            # Intentar con modelo de respaldo si está disponible
            if self.fallback_models:
                logger.info("Intentando con modelo de respaldo...")
                try:
                    fallback_model = next(iter(self.fallback_models.values()))
                    if file_path.suffix.lower() == '.pdf':
                        transcription = self.extract_pdf_content(str(file_path))
                        certificate_info = fallback_model.get_inference(transcription)
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        transcription = self.extract_docx_content(str(file_path))
                        certificate_info = fallback_model.get_inference(transcription)
                    elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
                        certificate_info = self.fallback_models.get("gemini", GeminiInferenceReport()).get_inference(path_file=str(file_path))
                        transcription = f"Imagen procesada por modelo de respaldo: {file_path.name}"
                    
                    logger.info("Éxito con modelo de respaldo")
                    return certificate_info, transcription
                except Exception as fallback_error:
                    logger.error(f"Error también en modelo de respaldo: {fallback_error}")
            
            raise

class ExcelReportManager:
    """Gestor para crear y manejar reportes Excel completos"""
    
    def __init__(self):
        self.headers = [
            "Nombre del Archivo",  # Nueva primera columna
            "Nombre del Certificado",
            "Nombre Completo", 
            "Identificación",
            "Institución",
            "Ciudad",
            "Fecha Emisión",
            "Fecha Expiración",
            "Estado",
            "Intensidad (hrs)",
            "Dirigido a",
            "Área",
            "Nivel",
            "Lineamientos",
            "Instructor",
            "NIT Institución",
            "Transcripción",
            "Mensaje Error"
        ]

    def create_excel_template(self, current_path: str, name_report: str = "reporte_certificaciones") -> Path:
        """Crea plantilla Excel con formato profesional incluyendo fecha de procesamiento"""
        try:
            excel_path = Path(current_path) / f'{name_report}.xlsx'
            
            if excel_path.exists():
                logger.info(f"Archivo existente encontrado: {excel_path}")
                return excel_path
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Certificaciones"
            
            # NUEVA CARACTERÍSTICA: Agregar fecha de procesamiento en la primera fila
            fecha_procesamiento = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            ws.merge_cells('A1:D1')  # Combinar celdas para la fecha
            fecha_cell = ws['A1']
            fecha_cell.value = f"📅 REPORTE GENERADO EL: {fecha_procesamiento}"
            fecha_cell.font = Font(bold=True, size=14, color="FFFFFF")
            fecha_cell.fill = PatternFill(start_color="1565C0", end_color="1565C0", fill_type="solid")
            fecha_cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Aplicar encabezados con formato en la fila 2
            for col, header in enumerate(self.headers, 1):
                cell = ws.cell(row=2, column=col)
                cell.value = header
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="2E86AB", end_color="2E86AB", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Ajustar anchos de columnas (agregando una más para el nombre del archivo)
            column_widths = [20, 25, 30, 15, 30, 15, 15, 15, 12, 12, 15, 15, 12, 20, 25, 15, 50, 30]
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[ws.cell(row=2, column=col).column_letter].width = width
            
            # Aplicar filtros automáticos (desde la fila 2)
            ws.auto_filter.ref = f"A2:{ws.cell(row=2, column=len(self.headers)).column_letter}2"
            
            wb.save(str(excel_path))
            logger.info(f"Plantilla Excel creada con fecha de procesamiento: {excel_path}")
            
            return excel_path
            
        except Exception as e:
            logger.error(f"Error creando plantilla Excel: {str(e)}")
            raise

    def calculate_status(self, expiration_date_str):
        """Calcula el estado del certificado basado en fecha de expiración"""
        if not expiration_date_str:
            return "Sin fecha"
        
        try:
            # Convertir fecha ISO a datetime
            exp_date = datetime.fromisoformat(expiration_date_str.replace('Z', '+00:00'))
            now = datetime.now(exp_date.tzinfo)
            
            if exp_date < now:
                return "VENCIDO"
            elif (exp_date - now).days <= 30:
                return "POR VENCER"
            else:
                return "VIGENTE"
                
        except Exception:
            return "Fecha inválida"

    def insert_certificate_data(self, excel_path: Path, data: dict, transcription: str, filename: str) -> None:
        """Inserta datos completos del certificado en Excel incluyendo nombre del archivo"""
        try:
            wb = load_workbook(str(excel_path))
            ws = wb.active
            
            # La primera fila es la fecha, los datos empiezan desde la fila 3
            last_row = ws.max_row + 1
            
            # Calcular estado
            status = self.calculate_status(data.get('expiration_date'))
            
            # Formatear fechas para Excel
            issue_date = self.format_date_for_excel(data.get('issue_date'))
            expiration_date = self.format_date_for_excel(data.get('expiration_date'))
            
            # NUEVA CARACTERÍSTICA: Insertar datos incluyendo nombre del archivo como primera columna
            row_data = [
                filename,  # Primera columna: nombre del archivo
                data.get('certificate_name', ''),
                data.get('participant_name', ''),
                data.get('identification', ''),
                data.get('institution', ''),
                data.get('city', ''),
                issue_date,
                expiration_date,
                status,
                data.get('hours', ''),
                data.get('target_audience', ''),
                data.get('specialization_area', ''),
                data.get('level', ''),
                data.get('guidelines', ''),
                data.get('instructor', ''),
                data.get('institution_nit', ''),
                transcription[:1000] + '...' if len(transcription) > 1000 else transcription,  # Limitar transcripción
                data.get('message_error', '')
            ]
            
            for col, value in enumerate(row_data, 1):
                ws.cell(row=last_row, column=col, value=value)
            
            # Aplicar formato condicional al estado (columna 9 ahora)
            status_cell = ws.cell(row=last_row, column=9)
            if status == "VENCIDO":
                status_cell.fill = PatternFill(start_color="FFCDD2", end_color="FFCDD2", fill_type="solid")
            elif status == "POR VENCER":
                status_cell.fill = PatternFill(start_color="FFF9C4", end_color="FFF9C4", fill_type="solid")
            elif status == "VIGENTE":
                status_cell.fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
            
            wb.save(str(excel_path))
            logger.info(f"Datos insertados en fila {last_row} para archivo: {filename}")
            
        except Exception as e:
            logger.error(f"Error insertando datos en Excel: {str(e)}")
            raise

    def format_date_for_excel(self, iso_date_str):
        """Convierte fecha ISO a formato legible para Excel"""
        if not iso_date_str:
            return ""
        
        try:
            date_obj = datetime.fromisoformat(iso_date_str.replace('Z', '+00:00'))
            return date_obj.strftime('%d/%m/%Y')
        except Exception:
            return iso_date_str

def process_single_file(file_path: str, excel_manager: ExcelReportManager, extractor: DocumentExtractorReport, 
                       excel_path: Path, folder_name: str) -> Tuple[bool, str]:
    """
    🆕 NUEVA FUNCIÓN: Procesa un archivo individual
    
    Returns:
        Tuple[bool, str]: (éxito, mensaje_error)
    """
    try:
        filename = os.path.basename(file_path)
        logger.info(f"📄 Procesando archivo: {filename}")
        
        # Extraer información y transcripción
        certificate_info, transcription = extractor.extract_content(file_path)
        
        # Insertar en Excel
        excel_manager.insert_certificate_data(excel_path, certificate_info, transcription, filename)
        
        logger.info(f"✅ Éxito: {filename}")
        return True, ""
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"❌ Error procesando {filename}: {error_msg}")
        
        # Insertar fila de error
        error_data = {
            'certificate_name': '',
            'message_error': error_msg
        }
        try:
            excel_manager.insert_certificate_data(excel_path, error_data, f"Error al procesar: {error_msg}", filename)
        except Exception as excel_error:
            logger.error(f"Error insertando error en Excel: {excel_error}")
        
        return False, error_msg

def process_multipage_pdf(pdf_path: str, extractor: DocumentExtractorReport, 
                         excel_manager: ExcelReportManager) -> Tuple[int, int]:
    """
    🆕 NUEVA FUNCIÓN PRINCIPAL: Maneja PDFs con múltiples páginas
    
    Args:
        pdf_path: Ruta del PDF multipágina
        extractor: Instancia del extractor
        excel_manager: Gestor de Excel
        
    Returns:
        Tuple[int, int]: (archivos_procesados_exitosamente, archivos_con_error)
    """
    logger.info("="*60)
    logger.info(f"🔄 PROCESANDO PDF MULTIPÁGINA: {os.path.basename(pdf_path)}")
    logger.info("="*60)
    
    # Obtener información del archivo
    pdf_path_obj = Path(pdf_path)
    base_name = pdf_path_obj.stem
    parent_dir = pdf_path_obj.parent
    
    # Crear directorio para las páginas separadas
    split_directory = parent_dir / f"{base_name}_paginas_separadas"
    
    # Separar el PDF en páginas individuales
    success, created_files = extractor.pdf_splitter.split_pdf_pages(pdf_path, str(split_directory))
    
    if not success or not created_files:
        logger.error(f"❌ No se pudieron separar las páginas del PDF: {pdf_path}")
        return 0, 1
    
    # Crear archivo Excel para este PDF multipágina
    excel_path = excel_manager.create_excel_template(str(split_directory), f"reporte_{base_name}")
    
    logger.info(f"📊 Creado archivo Excel: {excel_path}")
    
    # Procesar cada página separada
    successful_extractions = 0
    failed_extractions = 0
    
    for page_file in created_files:
        success, error_msg = process_single_file(
            page_file, excel_manager, extractor, excel_path, base_name
        )
        
        if success:
            successful_extractions += 1
        else:
            failed_extractions += 1
    
    # Mostrar resumen del PDF multipágina
    logger.info("="*60)
    logger.info(f"📋 RESUMEN PDF MULTIPÁGINA '{base_name}':")
    logger.info(f"   📄 Páginas separadas: {len(created_files)}")
    logger.info(f"   ✅ Procesadas exitosamente: {successful_extractions}")
    logger.info(f"   ❌ Con errores: {failed_extractions}")
    logger.info(f"   📊 Excel generado: {excel_path.name}")
    logger.info(f"   📂 Archivos en: {split_directory}")
    logger.info("="*60)
    
    return successful_extractions, failed_extractions

def process_certificates_directory():
    """🆕 FUNCIÓN PRINCIPAL MEJORADA: Procesa todos los certificados incluyendo manejo de PDFs multipágina"""
    
    logger.info("=" * 80)
    logger.info(f"🚀 INICIANDO PROCESAMIENTO CON MODELO: {DEFAULT_MODEL.upper()}")
    logger.info("=" * 80)
    logger.info("📋 CARACTERÍSTICAS IMPLEMENTADAS:")
    logger.info("   ✅ Primera fila: Fecha de procesamiento destacada")
    logger.info("   ✅ Primera columna: Nombre del archivo para trazabilidad")
    logger.info("   ✅ Estado de vigencia calculado desde fecha de procesamiento")
    logger.info("   🆕 NUEVA: Separación automática de PDFs multipágina")
    logger.info("   🆕 NUEVA: Procesamiento individual de cada página/certificado")
    logger.info("   🆕 NUEVA: Reportes independientes para PDFs multipágina")
    logger.info("=" * 80)
    
    # Configuración
    extractor = DocumentExtractorReport()
    excel_manager = ExcelReportManager()
    
    # Directorios
    directorio_actual = os.path.dirname(os.path.abspath(__file__))
    carpeta_certificates = os.path.join(directorio_actual, 'certificates')
    
    if not os.path.exists(carpeta_certificates):
        logger.error(f"El directorio {carpeta_certificates} no existe")
        return
    
    # Extensiones prohibidas
    extensiones_prohibidas = {ext.value for ext in ArchivoProhibido}
    
    # Contadores para estadísticas GLOBALES
    total_processed = 0
    successful_extractions = 0
    failed_extractions = 0
    multipage_pdfs_processed = 0
    
    # Procesar cada carpeta
    for raiz, dirs, archivos in os.walk(carpeta_certificates):
        nombre_carpeta = os.path.basename(raiz)
        
        # Filtrar archivos permitidos
        archivos_permitidos = [
            archivo for archivo in archivos 
            if os.path.splitext(archivo)[1].lower() not in extensiones_prohibidas
        ]
        
        if not archivos_permitidos:
            continue
            
        logger.info(f"📁 Procesando carpeta: {nombre_carpeta} ({len(archivos_permitidos)} archivos)")
        
        # Crear archivo Excel para archivos de página única en esta carpeta
        excel_path = excel_manager.create_excel_template(raiz, f"reporte_{nombre_carpeta}")
        
        # Procesar cada archivo
        for archivo in archivos_permitidos:
            path_file = os.path.join(raiz, archivo)
            total_processed += 1
            
            # 🆕 NUEVA LÓGICA: Verificar si es PDF multipágina
            if extractor.should_split_pdf(path_file):
                logger.info(f"🔄 PDF multipágina detectado: {archivo}")
                multipage_pdfs_processed += 1
                
                # Procesar PDF multipágina por separado
                mp_success, mp_failed = process_multipage_pdf(path_file, extractor, excel_manager)
                successful_extractions += mp_success
                failed_extractions += mp_failed
                
            else:
                # Procesar archivo normalmente (PDF de 1 página u otros formatos)
                success, error_msg = process_single_file(
                    path_file, excel_manager, extractor, excel_path, nombre_carpeta
                )
                
                if success:
                    successful_extractions += 1
                else:
                    failed_extractions += 1
    
    # Mostrar estadísticas finales MEJORADAS
    logger.info("=" * 80)
    logger.info("📊 RESUMEN FINAL DE PROCESAMIENTO:")
    logger.info(f"🤖 Modelo utilizado: {DEFAULT_MODEL.upper()}")
    logger.info(f"📝 Total de archivos procesados: {total_processed}")
    logger.info(f"📄 PDFs multipágina detectados: {multipage_pdfs_processed}")
    logger.info(f"✅ Extracciones exitosas: {successful_extractions}")
    logger.info(f"❌ Extracciones fallidas: {failed_extractions}")
    if total_processed > 0:
        logger.info(f"📈 Tasa de éxito: {(successful_extractions/total_processed*100):.1f}%")
    else:
        logger.info("📈 Tasa de éxito: N/A")
    logger.info("=" * 80)
    
    if multipage_pdfs_processed > 0:
        logger.info("🎯 NOTA: Los PDFs multipágina fueron separados automáticamente")
        logger.info("   y procesados como certificados individuales en carpetas separadas.")

if __name__ == "__main__":
    try:
        process_certificates_directory()
        logger.info("🎉 Procesamiento completado exitosamente")
    except Exception as e:
        logger.error(f"💥 Error crítico en el procesamiento: {str(e)}")
        raise